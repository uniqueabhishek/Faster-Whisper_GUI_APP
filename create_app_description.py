"""Script to create a Word document with application description."""

import sys
import os

# Configure stdout encoding for UTF-8 support
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')  # type: ignore[attr-defined]

from docx import Document  # noqa: E402
from docx.shared import Pt, RGBColor, Inches  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402

# Create document
doc = Document()

# Title
title = doc.add_heading('Faster-Whisper AI Transcriber', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Professional Audio-to-Text Solution')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Overview Section
doc.add_heading('Overview', 1)
doc.add_paragraph(
    'A powerful desktop application for converting audio and video files into accurate '
    'text transcriptions using state-of-the-art AI technology powered by OpenAI\'s Whisper model.'
)

doc.add_paragraph()

# Key Features Section
doc.add_heading('Key Features', 1)

# Modern User Interface
doc.add_heading('✓ Modern User Interface', 2)
ui_features = [
    'Sleek dark-themed GUI with intuitive sidebar navigation',
    'Seamless workflow between Preprocessing and Transcription views',
    'Drag-and-drop file support for effortless operation',
    'Real-time progress tracking and live logs',
    'Independent preprocessing windows for multi-tasking'
]
for feature in ui_features:
    doc.add_paragraph(feature, style='List Bullet')

# Advanced Audio Preprocessing
doc.add_heading('✓ Advanced Audio Preprocessing', 2)
preprocessing_features = [
    'Intelligent Noise Reduction - Remove background noise with customizable strength and noise floor settings',
    'Music Removal - Filter out background music while preserving speech clarity',
    'Audio Normalization - Automatically balance audio levels for consistent quality',
    'Voice Activity Detection (VAD) - Trim silence and detect speech segments automatically',
    'Format Conversion - Convert any audio format to optimized WAV with configurable sample rates (16kHz, 22kHz, 44.1kHz, 48kHz)',
    'Flexible Audio Settings - Choose mono/stereo output and bit depth (16/24/32-bit)'
]
for feature in preprocessing_features:
    doc.add_paragraph(feature, style='List Bullet')

# Powerful Transcription Engine
doc.add_heading('✓ Powerful Transcription Engine', 2)
transcription_features = [
    'Multi-Language Support - Auto-detect or manually select from multiple languages (English, Hindi, Japanese, Chinese, German, and more)',
    'Precision Control - Two analysis modes:',
    '  • Precise Analysis (beam size 5) for balanced speed and accuracy',
    '  • Deep Analysis (beam size 10) for maximum accuracy',
    'Advanced Options:',
    '  • Word-level timestamps for precise timing',
    '  • Language detection for automatic identification',
    '  • Custom initial prompts for context-aware transcription',
    '  • Batch processing for multiple files'
]
for feature in transcription_features:
    if feature.startswith('  •') or feature.startswith('  '):
        p = doc.add_paragraph(feature.strip(), style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    else:
        doc.add_paragraph(feature, style='List Bullet')

# Professional Output Formats
doc.add_heading('✓ Professional Output Formats', 2)
output_features = [
    'Text Files (.txt) - Clean, readable transcripts',
    'SRT Subtitles (.srt) - Perfect for video captioning',
    'VTT Subtitles (.vtt) - Web-compatible subtitle format',
    'JSON (.json) - Structured data with detailed metadata'
]
for feature in output_features:
    doc.add_paragraph(feature, style='List Bullet')

# Performance & Efficiency
doc.add_heading('✓ Performance & Efficiency', 2)
performance_features = [
    'GPU Acceleration - Automatic CUDA detection for faster processing',
    'Multi-threaded Architecture - Utilize all CPU cores for optimal performance',
    'Lazy Model Loading - Only loads AI models when needed to save memory',
    'Batch Processing - Process multiple files in queue automatically',
    'Smart Caching - Efficient resource management'
]
for feature in performance_features:
    doc.add_paragraph(feature, style='List Bullet')

# User-Friendly Features
doc.add_heading('✓ User-Friendly Features', 2)
user_features = [
    'Persistent Settings - Remembers your preferences between sessions',
    'Queue Management - Add, remove, and organize files easily',
    'Real-time Status Updates - Track progress with detailed logging',
    'Error Handling - Graceful error recovery with detailed error messages',
    'Output Organization - Automatic file naming and directory management',
    'One-Click Folder Access - Instantly open output directories'
]
for feature in user_features:
    doc.add_paragraph(feature, style='List Bullet')

# Format Compatibility
doc.add_heading('✓ Format Compatibility', 2)
format_features = [
    'Audio: MP3, WAV, M4A, FLAC, OGG',
    'Video: MP4, MKV, WEBM (extracts audio automatically)'
]
for feature in format_features:
    doc.add_paragraph(feature, style='List Bullet')

# Professional Quality
doc.add_heading('✓ Professional Quality', 2)
quality_features = [
    'Offline operation (no internet required after setup)',
    'Bundled VAD model for reliable speech detection',
    'Exception handling and crash recovery',
    'Detailed debug logging for troubleshooting',
    'Clean, maintainable codebase'
]
for feature in quality_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph()

# Ideal For Section
doc.add_heading('Ideal For', 1)
ideal_for = [
    'Content creators and podcasters',
    'Journalists and researchers',
    'Students and educators',
    'Business professionals',
    'Video producers and editors',
    'Accessibility services',
    'Anyone needing fast, accurate transcriptions'
]
for item in ideal_for:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# Whisper Model Capacities Section
doc.add_heading('Whisper Model Capacities', 1)
doc.add_paragraph(
    'The application supports all OpenAI Whisper model sizes, allowing you to choose the '
    'right balance between speed and accuracy for your needs:'
)

# Create a simple table-like structure with text
models_info = [
    ('tiny', '39 M parameters', '~1 GB RAM', 'Fastest, basic accuracy - Good for quick drafts'),
    ('base', '74 M parameters', '~1 GB RAM', 'Fast, improved accuracy - Good for simple audio'),
    ('small', '244 M parameters', '~2 GB RAM', 'Balanced speed/accuracy - Recommended for most users'),
    ('medium', '769 M parameters', '~5 GB RAM', 'High accuracy - Professional quality'),
    ('large-v1/v2', '1550 M parameters', '~10 GB RAM', 'Highest accuracy - Best for critical work'),
    ('large-v3', '1550 M parameters', '~10 GB RAM', 'Latest & most accurate - Cutting-edge performance')
]

for model, params, ram, description in models_info:
    p = doc.add_paragraph(style='List Bullet')
    # Model name in bold
    run = p.add_run(f'{model.upper()}: ')
    run.bold = True
    # Parameters
    p.add_run(f'{params} | {ram} | ')
    # Description
    p.add_run(description)

doc.add_paragraph()

# Performance Notes
doc.add_heading('Performance Notes:', 2)
perf_notes = [
    'GPU Acceleration: With CUDA-capable GPU, even large models run 5-10x faster',
    'CPU Performance: All models work on CPU, though larger models are slower',
    'Memory: Requirements decrease with GPU acceleration',
    'Recommended: "small" model for most users, "medium" or "large" for professional work',
    'Model Selection: Choose based on your hardware, accuracy needs, and speed requirements'
]
for note in perf_notes:
    doc.add_paragraph(note, style='List Bullet')

doc.add_paragraph()

# Technical Highlights Section
doc.add_heading('Technical Highlights', 1)
technical_features = [
    'Built with PyQt5 for cross-platform compatibility',
    'Powered by Faster-Whisper (optimized implementation of OpenAI Whisper)',
    'FFmpeg integration for robust audio processing',
    'Thread-safe architecture with concurrent processing',
    'Modular design for easy maintenance and extension'
]
for feature in technical_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer = doc.add_paragraph('Version 2.1.3')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_format = footer.runs[0]
footer_format.font.size = Pt(10)
footer_format.font.color.rgb = RGBColor(128, 128, 128)

developer = doc.add_paragraph('Developed by Abhishek\'s AI Labs')
developer.alignment = WD_ALIGN_PARAGRAPH.CENTER
developer_format = developer.runs[0]
developer_format.font.size = Pt(10)
developer_format.font.color.rgb = RGBColor(128, 128, 128)

# Save document
output_file = 'Faster-Whisper_AI_Transcriber_Description.docx'

# If file exists and is locked, try alternative name
if os.path.exists(output_file):
    try:
        # Try to delete if possible
        os.remove(output_file)
    except OSError:
        # If locked, use new name
        output_file = 'Faster-Whisper_AI_Transcriber_Description_v2.docx'

doc.save(output_file)
print(f"✓ Word document created successfully: {output_file}")
